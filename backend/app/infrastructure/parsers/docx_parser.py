"""DOCX parser using python-docx with style-level heading detection."""

from __future__ import annotations

import asyncio
from functools import partial

from docx import Document

from app.infrastructure.parsers.base import BaseParser, DocumentParseResult, ParsedSection
from app.infrastructure.parsers.section_detector import detect_heading_level


class DocxParser(BaseParser):
    def supports(self, filename: str) -> bool:
        return filename.lower().endswith(".docx")

    async def parse(self, file_path: str) -> DocumentParseResult:
        loop = asyncio.get_running_loop()
        doc = await loop.run_in_executor(None, partial(Document, file_path))
        all_text_parts: list[str] = []
        sections: list[ParsedSection] = []
        current_section = ParsedSection(heading="(preamble)")
        current_lines: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            all_text_parts.append(text)

            # Heading detection: docx style names or pattern matching
            style_name = (para.style.name or "").lower()
            is_heading = style_name.startswith("heading") or detect_heading_level(text) > 0

            if is_heading:
                current_section.content = "\n".join(current_lines).strip()
                if current_section.content or current_section.heading != "(preamble)":
                    sections.append(current_section)

                level = 1
                if style_name.startswith("heading"):
                    try:
                        level = int(style_name.replace("heading", "").strip())
                    except ValueError:
                        level = 1
                else:
                    level = detect_heading_level(text) or 1

                current_section = ParsedSection(heading=text, level=level)
                current_lines = []
            else:
                current_lines.append(text)

        # Final section
        current_section.content = "\n".join(current_lines).strip()
        if current_section.content:
            sections.append(current_section)

        # Also extract tables
        table_texts: list[str] = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_texts.append(" | ".join(cells))

        if table_texts:
            all_text_parts.append("\n[Table Data]\n" + "\n".join(table_texts))
            sections.append(
                ParsedSection(heading="Table Data", content="\n".join(table_texts), level=2)
            )

        return DocumentParseResult(
            raw_text="\n".join(all_text_parts),
            sections=sections,
            metadata={"paragraph_count": len(doc.paragraphs)},
        )
